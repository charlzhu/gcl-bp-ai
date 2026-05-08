#!/usr/bin/env python3
"""E2E QA Phase 0 数据资产准备脚本。

本脚本只负责附件解压、源数据画像和样例题清单生成，不计算标准答案，
也不读取或修改业务系统数据。所有输出均落在 ai/eval 目录下，便于复跑
和后续 Phase 1/Phase 2 消费。
"""

from __future__ import annotations

import argparse
import json
import re
import sys
import zipfile
from collections import Counter
from datetime import datetime
from pathlib import Path, PurePosixPath
from typing import Any

import pandas as pd
from docx import Document


LOGISTICS_ZIP_NAME = "23 年至 25 年物流源数据.zip"
BOM_ZIP_NAME = "BOM 源数据.zip"
QUESTION_DOCX_NAME = "物流和 bom样例题.docx"

LOGISTICS_KEYWORD_GROUPS: dict[str, list[str]] = {
    "date": ["日期", "发运日期", "发货日期", "出货日期", "月份"],
    "region": ["区域", "大区"],
    "province": ["省", "省份"],
    "city": ["市", "城市"],
    "origin": ["始发", "起运", "发货基地", "基地"],
    "customer": ["客户", "项目"],
    "product_spec": ["规格", "型号", "产品"],
    "transport_mode": ["运输方式", "运输类型", "承运方式"],
    "carrier": ["物流公司", "承运商", "供应商"],
    "shipment_count": ["发运件数", "实际发运", "发货件数", "组件数量", "数量"],
    "planned_count": ["计划发运", "计划件数"],
    "watt": ["瓦数", "功率", "mw", "MW", "兆瓦"],
    "pallet": ["托数", "托"],
    "vehicle": ["车次", "车辆", "车数", "车牌"],
    "fee": ["费用", "运费", "总费用", "金额"],
    "unit_price": ["单价", "元/瓦", "元/W", "元/车"],
}

BOM_KEYWORD_GROUPS: dict[str, list[str]] = {
    "order_no": ["订单", "订单号", "合同号", "项目号", "Sales Order"],
    "model": ["型号", "组件型号", "产品型号", "Model"],
    "material_code": ["物料编码", "材料编码", "Item", "Material"],
    "material_name": ["物料名称", "材料名称", "物料", "材料", "Description"],
    "spec_desc": ["规格", "规格描述", "描述", "物料描述", "Specification"],
    "quantity": ["用量", "数量", "Qty", "Quantity"],
    "unit": ["单位", "Unit"],
}

BOM_MATERIAL_TERMS = ["玻璃", "间隙贴膜", "焊带", "汇流条", "接线盒", "线盒"]

BOM_QUESTION_TERMS = [
    "bom",
    "bill of materials",
    "billofmaterials",
    "物料",
    "材料",
    "玻璃",
    "间隙贴膜",
    "焊带",
    "汇流条",
    "接线盒",
    "线盒",
    "规格描述",
    "规格对比",
    "NT10",
    "NT12",
]

LOGISTICS_QUESTION_TERMS = [
    "物流",
    "发运",
    "发货",
    "发往",
    "发出",
    "始发",
    "运输",
    "运量",
    "运费",
    "运价",
    "报价",
    "车次",
    "车辆",
    "车型",
    "托数",
    "区域",
    "省",
    "城市",
    "物流公司",
    "承运商",
    "公路",
    "铁路",
    "水路",
    "汽运",
    "铁运",
    "元/瓦",
    "成本",
    "费用",
    "客户",
    "线路",
    "基地",
    "路程",
    "MW",
    "mw",
    "任务",
    "招标",
    "询比价",
    "采购方式",
    "装车",
    "司机",
    "手机号",
    "身份证",
    "在途",
    "达标率",
    "特殊订单",
]


def repo_root_from_script() -> Path:
    """返回仓库根目录。

    参数：无。
    返回值：脚本所在路径向上回溯得到的仓库根目录。
    """

    return Path(__file__).resolve().parents[3]


def clean_cell(value: Any) -> str:
    """清洗 Excel 单元格文本。

    参数：
        value：pandas 读取到的原始单元格值。
    返回值：去除空值、全角空格和多余空白后的字符串。
    """

    if value is None:
        return ""
    if pd.isna(value):
        return ""
    text = str(value).replace("\u3000", " ").strip()
    if text.lower() in {"nan", "none"}:
        return ""
    return re.sub(r"\s+", " ", text)


def repair_zip_member_name(name: str) -> str:
    """修复 zip 内中文文件名的常见 mojibake。

    参数：
        name：zipfile 返回的成员名。
    返回值：优先返回可还原的 UTF-8 中文名，失败时返回原始成员名。
    """

    try:
        repaired = name.encode("cp437").decode("utf-8")
    except (UnicodeEncodeError, UnicodeDecodeError):
        return name
    chinese_count = sum("\u4e00" <= char <= "\u9fff" for char in repaired)
    original_noise = sum(char in name for char in "σΦ╣╛╜╡")
    if chinese_count > 0 or original_noise > 0:
        return repaired
    return name


def is_filtered_member(relative_name: str) -> bool:
    """判断 zip 成员是否属于 macOS 元数据或隐藏伴生文件。

    参数：
        relative_name：已修复后的 zip 相对路径。
    返回值：需要过滤时返回 True，否则返回 False。
    """

    parts = PurePosixPath(relative_name).parts
    return any(part == "__MACOSX" or part.startswith("._") for part in parts)


def safe_relative_path(relative_name: str) -> Path | None:
    """将 zip 成员名转换为安全的相对路径。

    参数：
        relative_name：zip 成员相对路径。
    返回值：安全 Path；如果路径为空或存在越界片段则返回 None。
    """

    parts = []
    for part in PurePosixPath(relative_name).parts:
        if part in {"", "."}:
            continue
        if part == "..":
            return None
        parts.append(part)
    if not parts:
        return None
    return Path(*parts)


def format_size(size: int) -> str:
    """格式化字节数。

    参数：
        size：字节数。
    返回值：便于报告阅读的容量字符串。
    """

    value = float(size)
    for unit in ["B", "KB", "MB", "GB"]:
        if value < 1024 or unit == "GB":
            return f"{value:.1f} {unit}" if unit != "B" else f"{int(value)} B"
        value /= 1024
    return f"{size} B"


def extract_zip(zip_path: Path, extract_root: Path) -> dict[str, Any]:
    """解压单个附件 zip 并过滤 macOS 元数据。

    参数：
        zip_path：附件 zip 路径。
        extract_root：统一解压根目录。
    返回值：包含提取文件、过滤文件和失败信息的 manifest。
    """

    target_root = extract_root / zip_path.stem
    target_root.mkdir(parents=True, exist_ok=True)
    manifest: dict[str, Any] = {
        "zip_name": zip_path.name,
        "zip_path": str(zip_path),
        "target_dir": str(target_root),
        "extracted_files": [],
        "filtered_members": [],
        "errors": [],
    }
    with zipfile.ZipFile(zip_path) as archive:
        for info in archive.infolist():
            repaired_name = repair_zip_member_name(info.filename)
            if is_filtered_member(repaired_name):
                manifest["filtered_members"].append(repaired_name)
                continue
            safe_path = safe_relative_path(repaired_name)
            if safe_path is None:
                manifest["filtered_members"].append(repaired_name)
                continue
            output_path = target_root / safe_path
            try:
                if info.is_dir() or repaired_name.endswith("/"):
                    output_path.mkdir(parents=True, exist_ok=True)
                    continue
                output_path.parent.mkdir(parents=True, exist_ok=True)
                with archive.open(info) as src, output_path.open("wb") as dst:
                    dst.write(src.read())
                manifest["extracted_files"].append(
                    {
                        "member": repaired_name,
                        "path": str(output_path),
                        "size": info.file_size,
                    }
                )
            except Exception as exc:  # pragma: no cover - 报告需要保留真实异常文本
                manifest["errors"].append(
                    {
                        "member": repaired_name,
                        "error": f"{type(exc).__name__}: {exc}",
                    }
                )
    return manifest


def detect_header_row(df: pd.DataFrame, keyword_groups: dict[str, list[str]]) -> int:
    """在工作表前若干行中识别最可能的表头行。

    参数：
        df：按无表头方式读取的工作表。
        keyword_groups：用于当前领域的字段关键词集合。
    返回值：表头行的 0 基索引。
    """

    keywords = [keyword.lower() for values in keyword_groups.values() for keyword in values]
    max_scan_rows = min(len(df), 40)
    best_row = 0
    best_score = -1.0
    for row_index in range(max_scan_rows):
        cells = [clean_cell(value) for value in df.iloc[row_index].tolist()]
        non_empty = [cell for cell in cells if cell]
        if not non_empty:
            continue
        keyword_hits = 0
        for cell in non_empty:
            lower_cell = cell.lower()
            if any(keyword in lower_cell for keyword in keywords):
                keyword_hits += 1
        short_text_count = sum(1 for cell in non_empty if len(cell) <= 40)
        duplicate_penalty = len(non_empty) - len(set(non_empty))
        score = keyword_hits * 10 + min(len(non_empty), 30) * 1.5 + short_text_count * 0.2
        score -= duplicate_penalty * 0.5
        if len(non_empty) < 2:
            score -= 8
        if score > best_score:
            best_score = score
            best_row = row_index
    return best_row


def build_fields(df: pd.DataFrame, header_row: int) -> list[str]:
    """从表头行构造字段列表。

    参数：
        df：无表头工作表数据。
        header_row：表头行索引。
    返回值：去重后的字段名列表，空字段以“未命名列”补齐。
    """

    raw_fields = [clean_cell(value) for value in df.iloc[header_row].tolist()]
    fields: list[str] = []
    seen: Counter[str] = Counter()
    for index, field in enumerate(raw_fields, start=1):
        normalized = field or f"未命名列{index}"
        seen[normalized] += 1
        if seen[normalized] > 1:
            normalized = f"{normalized}_{seen[normalized]}"
        fields.append(normalized)
    while fields and fields[-1].startswith("未命名列"):
        fields.pop()
    return fields


def count_data_rows(df: pd.DataFrame, header_row: int) -> int:
    """统计表头之后的非空数据行数。

    参数：
        df：无表头工作表数据。
        header_row：表头行索引。
    返回值：非空数据行数，不包含表头行。
    """

    count = 0
    for row_index in range(header_row + 1, len(df)):
        row_values = [clean_cell(value) for value in df.iloc[row_index].tolist()]
        if any(row_values):
            count += 1
    return count


def match_key_fields(fields: list[str], groups: dict[str, list[str]]) -> list[dict[str, str]]:
    """识别字段列表中的关键业务字段。

    参数：
        fields：字段名列表。
        groups：领域关键词分组。
    返回值：匹配到的业务字段分组、字段名和命中关键词。
    """

    matches: list[dict[str, str]] = []
    for group, keywords in groups.items():
        for field in fields:
            lower_field = field.lower()
            for keyword in keywords:
                if keyword.lower() in lower_field:
                    matches.append({"group": group, "field": field, "keyword": keyword})
                    break
    return matches


def guess_years_from_text(text: str) -> list[int]:
    """从文件名、sheet 名或问题文本中识别年份。

    参数：
        text：待分析文本。
    返回值：升序年份列表。
    """

    years = {int(year) for year in re.findall(r"20\d{2}", text)}
    for start, end in re.findall(r"(20\d{2})\s*[-~至到－—–]\s*(20\d{2})", text):
        start_year = int(start)
        end_year = int(end)
        if start_year <= end_year and end_year - start_year <= 10:
            years.update(range(start_year, end_year + 1))
    for start, end in re.findall(r"(?<!\d)0?(23|24|25|26)\s*[-~至到－—–]\s*0?(23|24|25|26)\s*年", text):
        start_year = 2000 + int(start)
        end_year = 2000 + int(end)
        if start_year <= end_year and end_year - start_year <= 10:
            years.update(range(start_year, end_year + 1))
    for short_year in re.findall(r"(?<!\d)0?(23|24|25|26)\s*年", text):
        years.add(2000 + int(short_year))
    return sorted(years)


def scan_material_hits(df: pd.DataFrame) -> dict[str, int]:
    """扫描 BOM 工作表中关键物料词出现次数。

    参数：
        df：无表头工作表数据。
    返回值：关键物料词到出现次数的映射。
    """

    values = [clean_cell(value) for value in df.to_numpy().ravel()]
    joined_text = "\n".join(value for value in values if value)
    return {term: joined_text.count(term) for term in BOM_MATERIAL_TERMS if joined_text.count(term) > 0}


def excel_engine_for(path: Path) -> str:
    """根据扩展名选择 Excel 读取引擎。

    参数：
        path：Excel 文件路径。
    返回值：pandas 可用的读取引擎名称。
    """

    return "xlrd" if path.suffix.lower() == ".xls" else "openpyxl"


def profile_excel_file(path: Path, domain: str) -> dict[str, Any]:
    """生成单个 Excel 文件画像。

    参数：
        path：Excel 或 xls 文件路径。
        domain：logistics 或 bom，用于选择字段识别规则。
    返回值：包含 sheet、字段、行数、年份和关键字段的画像字典。
    """

    groups = LOGISTICS_KEYWORD_GROUPS if domain == "logistics" else BOM_KEYWORD_GROUPS
    engine = excel_engine_for(path)
    profile: dict[str, Any] = {
        "path": str(path),
        "file_name": path.name,
        "domain": domain,
        "size": path.stat().st_size,
        "year_guess": guess_years_from_text(path.name),
        "sheet_count": 0,
        "total_data_rows": 0,
        "sheets": [],
        "errors": [],
    }
    if domain == "bom":
        profile.update(guess_bom_identity(path.name))
    try:
        excel_file = pd.ExcelFile(path, engine=engine)
        profile["sheet_count"] = len(excel_file.sheet_names)
        for sheet_name in excel_file.sheet_names:
            try:
                df = pd.read_excel(
                    path,
                    sheet_name=sheet_name,
                    header=None,
                    dtype=object,
                    engine=engine,
                )
                df = df.dropna(how="all").dropna(axis=1, how="all")
                if df.empty:
                    sheet_profile = {
                        "sheet_name": sheet_name,
                        "header_row": None,
                        "row_count": 0,
                        "field_count": 0,
                        "fields": [],
                        "key_fields": [],
                        "year_guess": guess_years_from_text(sheet_name),
                        "material_hits": {},
                    }
                else:
                    header_row = detect_header_row(df, groups)
                    fields = build_fields(df, header_row)
                    row_count = count_data_rows(df, header_row)
                    profile["total_data_rows"] += row_count
                    sheet_profile = {
                        "sheet_name": sheet_name,
                        "header_row": header_row + 1,
                        "row_count": row_count,
                        "field_count": len(fields),
                        "fields": fields,
                        "key_fields": match_key_fields(fields, groups),
                        "year_guess": guess_years_from_text(f"{path.name} {sheet_name}"),
                        "material_hits": scan_material_hits(df) if domain == "bom" else {},
                    }
                profile["sheets"].append(sheet_profile)
            except Exception as exc:  # pragma: no cover - 解析失败需进入报告
                profile["errors"].append(
                    {
                        "sheet_name": sheet_name,
                        "error": f"{type(exc).__name__}: {exc}",
                    }
                )
    except Exception as exc:  # pragma: no cover - 解析失败需进入报告
        profile["errors"].append({"sheet_name": None, "error": f"{type(exc).__name__}: {exc}"})
    return profile


def guess_bom_identity(file_name: str) -> dict[str, Any]:
    """从 BOM 文件名中猜测订单号和产品型号。

    参数：
        file_name：BOM 文件名。
    返回值：包含订单号、型号和版本片段的猜测结果。
    """

    order_numbers = sorted(set(re.findall(r"20\d{2}-\d{5}", file_name)))
    model_guess = ""
    model_match = re.match(r"([A-Za-z]{1,4}\d+[A-Za-z0-9R/]*)\(", file_name)
    if model_match:
        model_guess = model_match.group(1)
    version_match = re.search(r"Bill\s*of\s*materials[- ]*([A-Z])", file_name, flags=re.I)
    if not version_match:
        version_match = re.search(r"Billofmaterials[- ]*([A-Z])", file_name, flags=re.I)
    return {
        "order_no_guess": order_numbers,
        "model_guess": model_guess,
        "bom_version_guess": version_match.group(1).upper() if version_match else "",
    }


def guess_domain(question: str) -> str:
    """判断样例题更接近物流还是 BOM。

    参数：
        question：样例题文本。
    返回值：bom、logistics 或 unknown。
    """

    lower_question = question.lower()
    bom_score = sum(1 for term in BOM_QUESTION_TERMS if term.lower() in lower_question)
    logistics_score = sum(1 for term in LOGISTICS_QUESTION_TERMS if term.lower() in lower_question)
    if bom_score >= 2 and logistics_score == 0:
        return "bom"
    if bom_score >= 2 and "发运" not in question and "物流" not in question:
        return "bom"
    if logistics_score > 0:
        return "logistics"
    if "订单" in question and ("规格" in question or "材料" in question):
        return "bom"
    return "unknown"


def guess_category(question: str, domain: str) -> str:
    """根据问题文本猜测 E2E 分类。

    参数：
        question：样例题文本。
        domain：领域猜测结果。
    返回值：可供后续 Phase 使用的粗分类。
    """

    if domain == "bom":
        if any(term in question for term in ["对比", "不一致", "不一样", "比较"]):
            return "bom_material_compare"
        if any(term in question for term in ["表格", "EXCEL", "excel", "统计出来"]):
            return "bom_material_table"
        return "bom_material_spec"
    if domain == "logistics":
        if any(term in question for term in ["司机", "手机号", "身份证"]):
            return "logistics_driver_consistency"
        if any(term in question for term in ["采购方式", "询比价", "招标任务", "招标"]):
            return "logistics_procurement_task"
        if any(term in question for term in ["是不是", "看看", "异常", "风险", "特殊", "最近", "当前"]):
            return "logistics_ambiguous_or_current"
        if any(term in question for term in ["前5", "前五", "排名"]):
            return "logistics_topn"
        if "平均路程" in question or "路程" in question:
            return "logistics_distance"
        if "达标率" in question and any(term in question for term in ["均值", "中位数"]):
            return "logistics_rate_statistics"
        if "偏差率" in question:
            return "logistics_plan_actual_variance"
        if "平均每车装载托数" in question:
            return "logistics_loading_efficiency"
        if any(term in question for term in ["平均单价/车", "运价", "报价"]):
            return "logistics_company_unit_price"
        if "平均元/瓦" in question or ("成本" in question and "排序" in question):
            return "logistics_cost_sort"
        if "车次" in question or "车辆数" in question:
            return "logistics_vehicle_count"
        if "总瓦数" in question or "运量" in question:
            return "logistics_shipment_watt"
        if "总费用" in question:
            return "logistics_total_fee"
        if "运输方式" in question:
            return "logistics_transport_mode_count"
        if "总发运件数" in question or "总件数" in question or "发运记录数" in question:
            return "logistics_count"
        return "logistics_other"
    return "unknown"


def needs_db_2026(question: str, domain: str, years: list[int]) -> bool:
    """判断问题是否可能依赖 2026+ MySQL 数据。

    参数：
        question：样例题文本。
        domain：领域猜测结果。
        years：问题中的年份猜测。
    返回值：物流 2026+ 或当前态问题返回 True；BOM 附件题返回 False。
    """

    if domain != "logistics":
        return False
    if any(year >= 2026 for year in years):
        return True
    return any(term in question for term in ["当前", "最近", "在途"])


def guess_question_years(question: str, domain: str) -> list[int]:
    """识别样例题年份。

    参数：
        question：样例题文本。
        domain：领域猜测结果。
    返回值：年份列表；历史物流且无显式年份时返回 2023-2025。
    """

    years = guess_years_from_text(question)
    if not years and domain == "logistics" and "历史" in question:
        return [2023, 2024, 2025]
    if not years and domain == "logistics" and any(term in question for term in ["当前", "最近", "在途"]):
        return [2026]
    return years


def parse_question_docx(docx_path: Path) -> list[dict[str, Any]]:
    """解析样例题 docx 为 JSONL 记录。

    参数：
        docx_path：样例题 docx 路径。
    返回值：问题记录列表，每条包含 question_id、question、seq 和猜测字段。
    """

    document = Document(docx_path)
    text_lines: list[str] = []
    for paragraph in document.paragraphs:
        for line in paragraph.text.splitlines():
            line = line.strip()
            if line:
                text_lines.append(line)
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                for line in cell.text.splitlines():
                    line = line.strip()
                    if line:
                        text_lines.append(line)

    questions: list[dict[str, Any]] = []
    seen_seq: set[int] = set()
    pattern = re.compile(r"^\s*(\d{1,5})[\.\u3001\uff0e\uff09\)]\s*(.+?)\s*$")
    for line in text_lines:
        match = pattern.match(line)
        if not match:
            continue
        seq = int(match.group(1))
        question = match.group(2).strip()
        if not question or seq in seen_seq:
            continue
        seen_seq.add(seq)
        domain = guess_domain(question)
        years = guess_question_years(question, domain)
        questions.append(
            {
                "question_id": f"Q{seq:04d}",
                "question": question,
                "seq": seq,
                "domain_guess": domain,
                "year_guess": years,
                "needs_db_2026_guess": needs_db_2026(question, domain, years),
                "category_guess": guess_category(question, domain),
            }
        )
    questions.sort(key=lambda item: item["seq"])
    return questions


def write_jsonl(path: Path, records: list[dict[str, Any]]) -> None:
    """写入 JSONL 文件。

    参数：
        path：输出文件路径。
        records：待写入记录列表。
    返回值：无。
    """

    path.parent.mkdir(parents=True, exist_ok=True)
    with path.open("w", encoding="utf-8") as output:
        for record in records:
            output.write(json.dumps(record, ensure_ascii=False, separators=(",", ":")) + "\n")


def summarize_question_stats(questions: list[dict[str, Any]]) -> dict[str, Any]:
    """汇总样例题统计信息。

    参数：
        questions：样例题记录。
    返回值：领域、年份、分类和 2026 DB 猜测统计。
    """

    domain_counter = Counter(item["domain_guess"] for item in questions)
    category_counter = Counter(item["category_guess"] for item in questions)
    year_counter: Counter[str] = Counter()
    for item in questions:
        years = item["year_guess"]
        if years:
            for year in years:
                year_counter[str(year)] += 1
        else:
            year_counter["未识别"] += 1
    return {
        "total": len(questions),
        "domain": dict(sorted(domain_counter.items())),
        "category": dict(category_counter.most_common()),
        "year": dict(sorted(year_counter.items())),
        "needs_db_2026": sum(1 for item in questions if item["needs_db_2026_guess"]),
    }


def collect_excel_profiles(extract_root: Path) -> tuple[list[dict[str, Any]], list[dict[str, Any]]]:
    """收集物流与 BOM Excel 画像。

    参数：
        extract_root：附件解压根目录。
    返回值：物流画像列表和 BOM 画像列表。
    """

    logistics_root = extract_root / Path(LOGISTICS_ZIP_NAME).stem
    bom_root = extract_root / Path(BOM_ZIP_NAME).stem
    logistics_files = sorted(
        [
            path
            for path in logistics_root.rglob("*")
            if path.is_file() and path.suffix.lower() in {".xls", ".xlsx"}
        ]
    )
    bom_files = sorted(
        [
            path
            for path in bom_root.rglob("*")
            if path.is_file() and path.suffix.lower() in {".xls", ".xlsx"}
        ]
    )
    logistics_profiles = [profile_excel_file(path, "logistics") for path in logistics_files]
    bom_profiles = [profile_excel_file(path, "bom") for path in bom_files]
    return logistics_profiles, bom_profiles


def flatten_errors(profiles: list[dict[str, Any]]) -> list[dict[str, Any]]:
    """提取文件画像中的解析失败项。

    参数：
        profiles：Excel 文件画像列表。
    返回值：失败项列表。
    """

    errors: list[dict[str, Any]] = []
    for profile in profiles:
        for error in profile.get("errors", []):
            errors.append(
                {
                    "file": profile["file_name"],
                    "sheet": error.get("sheet_name"),
                    "error": error.get("error"),
                }
            )
    return errors


def brief_fields(fields: list[str], limit: int = 12) -> str:
    """生成字段预览文本。

    参数：
        fields：字段列表。
        limit：最多展示字段数。
    返回值：逗号拼接的字段预览。
    """

    if not fields:
        return "-"
    preview = fields[:limit]
    suffix = f" 等 {len(fields)} 个字段" if len(fields) > limit else ""
    return "、".join(preview) + suffix


def brief_key_fields(sheet: dict[str, Any], limit: int = 8) -> str:
    """生成关键字段摘要。

    参数：
        sheet：sheet 画像。
        limit：最多展示匹配数。
    返回值：关键字段摘要文本。
    """

    matches = sheet.get("key_fields") or []
    if not matches:
        return "-"
    values = [f"{item['group']}={item['field']}" for item in matches[:limit]]
    if len(matches) > limit:
        values.append(f"等 {len(matches)} 项")
    return "；".join(values)


def write_report(
    path: Path,
    attachment_manifest: list[dict[str, Any]],
    logistics_profiles: list[dict[str, Any]],
    bom_profiles: list[dict[str, Any]],
    question_stats: dict[str, Any],
    questions: list[dict[str, Any]],
    output_paths: dict[str, str],
) -> None:
    """写入 Phase 0 数据画像报告。

    参数：
        path：报告输出路径。
        attachment_manifest：附件和解压 manifest。
        logistics_profiles：物流 Excel 画像。
        bom_profiles：BOM xls 画像。
        question_stats：样例题统计。
        questions：样例题明细。
        output_paths：关键输出路径。
    返回值：无。
    """

    lines: list[str] = []
    lines.append("# E2E QA Phase 0 数据资产画像报告")
    lines.append("")
    lines.append(f"- 生成时间：{datetime.now().isoformat(timespec='seconds')}")
    lines.append("- 执行边界：仅准备数据资产，不计算标准答案，不访问数据库，不修改业务代码。")
    lines.append(f"- 解压目录：`{output_paths['extract_dir']}`")
    lines.append(f"- 样例题 JSONL：`{output_paths['sample_questions']}`")
    lines.append("")

    lines.append("## 附件清单")
    lines.append("")
    lines.append("| 附件 | 类型 | 大小 | 有效提取文件数 | 过滤项数 | 解析/解压错误 |")
    lines.append("| --- | --- | ---: | ---: | ---: | ---: |")
    for item in attachment_manifest:
        lines.append(
            "| {name} | {type} | {size} | {extracted} | {filtered} | {errors} |".format(
                name=item["name"],
                type=item["type"],
                size=format_size(item["size"]),
                extracted=item.get("extracted_count", 0),
                filtered=item.get("filtered_count", 0),
                errors=item.get("error_count", 0),
            )
        )
    lines.append("")
    lines.append("过滤规则：跳过路径中包含 `__MACOSX` 的成员，以及文件名以 `._` 开头的 macOS 资源分叉文件。")
    lines.append("")

    total_logistics_sheets = sum(profile["sheet_count"] for profile in logistics_profiles)
    total_logistics_rows = sum(profile["total_data_rows"] for profile in logistics_profiles)
    lines.append("## 物流 Excel 文件画像")
    lines.append("")
    lines.append(
        f"- 文件数：{len(logistics_profiles)}；sheet 数：{total_logistics_sheets}；非空数据行数合计：{total_logistics_rows}"
    )
    lines.append("")
    lines.append("| 文件 | 年份猜测 | sheet 数 | 数据行合计 | 解析错误 |")
    lines.append("| --- | --- | ---: | ---: | ---: |")
    for profile in logistics_profiles:
        lines.append(
            f"| {profile['file_name']} | {profile.get('year_guess') or '-'} | {profile['sheet_count']} | {profile['total_data_rows']} | {len(profile.get('errors', []))} |"
        )
    lines.append("")
    for profile in logistics_profiles:
        lines.append(f"### {profile['file_name']}")
        lines.append("")
        lines.append("| Sheet | 表头行 | 数据行 | 字段数 | 年份猜测 | 关键字段 | 字段预览 |")
        lines.append("| --- | ---: | ---: | ---: | --- | --- | --- |")
        for sheet in profile["sheets"]:
            lines.append(
                "| {sheet} | {header} | {rows} | {field_count} | {years} | {keys} | {fields} |".format(
                    sheet=sheet["sheet_name"],
                    header=sheet["header_row"] or "-",
                    rows=sheet["row_count"],
                    field_count=sheet["field_count"],
                    years=sheet.get("year_guess") or "-",
                    keys=brief_key_fields(sheet),
                    fields=brief_fields(sheet.get("fields", [])),
                )
            )
        lines.append("")

    total_bom_sheets = sum(profile["sheet_count"] for profile in bom_profiles)
    total_bom_rows = sum(profile["total_data_rows"] for profile in bom_profiles)
    lines.append("## BOM xls 文件画像")
    lines.append("")
    lines.append(f"- 文件数：{len(bom_profiles)}；sheet 数：{total_bom_sheets}；非空数据行数合计：{total_bom_rows}")
    lines.append("")
    lines.append("| 文件 | 订单号猜测 | 型号猜测 | 版本 | sheet 数 | 数据行合计 | 关键物料命中 | 解析错误 |")
    lines.append("| --- | --- | --- | --- | ---: | ---: | --- | ---: |")
    for profile in bom_profiles:
        material_hits = Counter()
        for sheet in profile.get("sheets", []):
            material_hits.update(sheet.get("material_hits", {}))
        material_text = "；".join(f"{key}:{value}" for key, value in material_hits.items()) or "-"
        lines.append(
            "| {file} | {order} | {model} | {version} | {sheets} | {rows} | {materials} | {errors} |".format(
                file=profile["file_name"],
                order="、".join(profile.get("order_no_guess", [])) or "-",
                model=profile.get("model_guess") or "-",
                version=profile.get("bom_version_guess") or "-",
                sheets=profile["sheet_count"],
                rows=profile["total_data_rows"],
                materials=material_text,
                errors=len(profile.get("errors", [])),
            )
        )
    lines.append("")
    lines.append("### BOM 字段与关键物料可能字段")
    lines.append("")
    for profile in bom_profiles:
        lines.append(f"- `{profile['file_name']}`")
        for sheet in profile.get("sheets", []):
            material_hits = sheet.get("material_hits") or {}
            material_text = "；".join(f"{key}:{value}" for key, value in material_hits.items()) or "-"
            lines.append(
                f"  - Sheet `{sheet['sheet_name']}`：字段 {sheet['field_count']} 个，数据行 {sheet['row_count']}，可能字段：{brief_key_fields(sheet)}，关键物料命中：{material_text}"
            )
    lines.append("")

    lines.append("## 样例题统计")
    lines.append("")
    lines.append(f"- 样例题总数：{question_stats['total']}")
    lines.append(f"- 领域分布：`{json.dumps(question_stats['domain'], ensure_ascii=False)}`")
    lines.append(f"- 年份分布：`{json.dumps(question_stats['year'], ensure_ascii=False)}`")
    lines.append(f"- 可能需要 2026 MySQL 的物流题：{question_stats['needs_db_2026']}")
    lines.append("")
    lines.append("| 分类 | 数量 |")
    lines.append("| --- | ---: |")
    for category, count in question_stats["category"].items():
        lines.append(f"| {category} | {count} |")
    lines.append("")
    if questions:
        lines.append("### 样例题首尾检查")
        lines.append("")
        head_tail = questions[:5] + ([{"question_id": "...", "question": "..."}] if len(questions) > 10 else []) + questions[-5:]
        for item in head_tail:
            lines.append(f"- `{item.get('question_id')}` {item.get('question')}")
        lines.append("")

    errors = flatten_errors(logistics_profiles) + flatten_errors(bom_profiles)
    zip_errors = []
    for item in attachment_manifest:
        zip_errors.extend(item.get("errors", []))
    lines.append("## 解析失败")
    lines.append("")
    if not errors and not zip_errors:
        lines.append("- 未发现 Excel sheet 解析失败或 zip 解压失败。")
    else:
        for error in zip_errors:
            lines.append(f"- zip 成员 `{error.get('member')}`：{error.get('error')}")
        for error in errors:
            lines.append(f"- 文件 `{error['file']}` / sheet `{error.get('sheet')}`：{error['error']}")
    lines.append("")

    path.parent.mkdir(parents=True, exist_ok=True)
    path.write_text("\n".join(lines), encoding="utf-8")


def build_attachment_manifest(attachment_dir: Path, extraction_manifests: list[dict[str, Any]]) -> list[dict[str, Any]]:
    """汇总附件清单与解压结果。

    参数：
        attachment_dir：附件目录。
        extraction_manifests：zip 解压 manifest 列表。
    返回值：报告用附件清单。
    """

    manifest_by_name = {Path(item["zip_name"]).name: item for item in extraction_manifests}
    attachments: list[dict[str, Any]] = []
    for path in sorted(attachment_dir.iterdir()):
        if not path.is_file():
            continue
        item: dict[str, Any] = {
            "name": path.name,
            "path": str(path),
            "type": path.suffix.lower().lstrip(".") or "unknown",
            "size": path.stat().st_size,
            "extracted_count": 0,
            "filtered_count": 0,
            "error_count": 0,
            "errors": [],
        }
        if path.name in manifest_by_name:
            zip_manifest = manifest_by_name[path.name]
            item["extracted_count"] = len(zip_manifest.get("extracted_files", []))
            item["filtered_count"] = len(zip_manifest.get("filtered_members", []))
            item["error_count"] = len(zip_manifest.get("errors", []))
            item["errors"] = zip_manifest.get("errors", [])
        attachments.append(item)
    return attachments


def run_phase0(repo_root: Path) -> dict[str, Any]:
    """执行 Phase 0 全流程。

    参数：
        repo_root：仓库根目录。
    返回值：包含输出路径和统计信息的执行摘要。
    """

    attachment_dir = repo_root / "ai" / "inbox" / "attachments"
    eval_dir = repo_root / "ai" / "eval"
    workdir = eval_dir / "workdir"
    extract_dir = workdir / "attachments_extracted"
    sample_questions_path = eval_dir / "sample_questions.jsonl"
    report_path = eval_dir / "data_profile_report.md"
    profile_json_path = workdir / "phase0_data_profile.json"

    required_files = [
        attachment_dir / LOGISTICS_ZIP_NAME,
        attachment_dir / BOM_ZIP_NAME,
        attachment_dir / QUESTION_DOCX_NAME,
    ]
    missing = [str(path) for path in required_files if not path.exists()]
    if missing:
        raise FileNotFoundError(f"缺少必要附件: {missing}")

    extract_dir.mkdir(parents=True, exist_ok=True)
    extraction_manifests = [
        extract_zip(attachment_dir / LOGISTICS_ZIP_NAME, extract_dir),
        extract_zip(attachment_dir / BOM_ZIP_NAME, extract_dir),
    ]
    attachment_manifest = build_attachment_manifest(attachment_dir, extraction_manifests)
    logistics_profiles, bom_profiles = collect_excel_profiles(extract_dir)
    questions = parse_question_docx(attachment_dir / QUESTION_DOCX_NAME)
    question_stats = summarize_question_stats(questions)

    write_jsonl(sample_questions_path, questions)
    output_paths = {
        "extract_dir": str(extract_dir),
        "sample_questions": str(sample_questions_path),
        "report": str(report_path),
        "profile_json": str(profile_json_path),
    }
    write_report(
        report_path,
        attachment_manifest,
        logistics_profiles,
        bom_profiles,
        question_stats,
        questions,
        output_paths,
    )
    profile_json_path.parent.mkdir(parents=True, exist_ok=True)
    profile_payload = {
        "generated_at": datetime.now().isoformat(timespec="seconds"),
        "attachments": attachment_manifest,
        "extraction": extraction_manifests,
        "logistics_profiles": logistics_profiles,
        "bom_profiles": bom_profiles,
        "question_stats": question_stats,
        "outputs": output_paths,
    }
    profile_json_path.write_text(json.dumps(profile_payload, ensure_ascii=False, indent=2), encoding="utf-8")

    return {
        "outputs": output_paths,
        "attachment_count": len(attachment_manifest),
        "logistics_file_count": len(logistics_profiles),
        "logistics_sheet_count": sum(item["sheet_count"] for item in logistics_profiles),
        "logistics_row_count": sum(item["total_data_rows"] for item in logistics_profiles),
        "bom_file_count": len(bom_profiles),
        "bom_sheet_count": sum(item["sheet_count"] for item in bom_profiles),
        "bom_row_count": sum(item["total_data_rows"] for item in bom_profiles),
        "sample_question_count": len(questions),
        "question_stats": question_stats,
        "parse_error_count": len(flatten_errors(logistics_profiles) + flatten_errors(bom_profiles))
        + sum(len(item.get("errors", [])) for item in extraction_manifests),
        "filtered_member_count": sum(len(item.get("filtered_members", [])) for item in extraction_manifests),
    }


def parse_args() -> argparse.Namespace:
    """解析命令行参数。

    参数：无。
    返回值：argparse 解析结果。
    """

    parser = argparse.ArgumentParser(description="E2E QA Phase 0 数据资产准备")
    parser.add_argument(
        "--repo-root",
        type=Path,
        default=repo_root_from_script(),
        help="仓库根目录，默认按脚本位置自动推断。",
    )
    return parser.parse_args()


def main() -> int:
    """脚本入口。

    参数：无。
    返回值：进程退出码，0 表示成功。
    """

    args = parse_args()
    summary = run_phase0(args.repo_root.resolve())
    print(json.dumps(summary, ensure_ascii=False, indent=2))
    return 0


if __name__ == "__main__":
    sys.exit(main())
