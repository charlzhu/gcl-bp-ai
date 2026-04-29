from __future__ import annotations

import argparse
import json
import re
import sys
from collections import Counter, defaultdict
from dataclasses import asdict, dataclass
from pathlib import Path
from typing import Any

from docx import Document

PROJECT_ROOT = Path(__file__).resolve().parents[1]
if str(PROJECT_ROOT) not in sys.path:
    sys.path.insert(0, str(PROJECT_ROOT))

from backend.app.domains.logistics.services.data_qa_planner import LogisticsDataQaPlanner


SUPPORTED_QUERY_KEY_DESCRIPTIONS = {
    "hist_total_fee_city_rank": "历史台账中按省份过滤的城市总费用排名",
    "hist_avg_fee_by_month": "历史台账中指定始发地/省份/车型的月均运费",
    "hist_avg_fee_per_watt_by_transport": "历史台账中按区域统计运输方式平均元/瓦",
    "hist_extra_fee_ratio_peak_month": "历史台账中年度额外费用占比最高月份",
    "hist_total_fee_by_origin_and_carrier": "历史台账中指定始发地和承运商的总运费",
    "hist_trip_count_by_region": "历史台账中按区域汇总车次",
    "hist_quantity_by_region": "历史台账中按区域汇总发运件数",
    "hist_customer_mw": "历史台账中指定客户/项目的发运量 MW",
    "hist_vehicle_type_trip_count": "历史台账中指定车型的发运车次",
    "hist_multi_origin_customers": "历史台账中同一客户多个始发地发货统计",
    "hist_plan_actual_deviation": "历史台账中计划件数与实际件数偏差率",
    "sys_mw_and_trip_count": "2026 系统侧按 pickup_date 统计发运量和车次",
    "sys_signedfor_rate_by_carrier": "2026 系统侧承运商 SIGNEDFOR 签收率前后十",
    "sys_companies_without_tasks": "2026 已建档承运商中无任务记录公司",
    "sys_special_total_fee": "2026 特殊业务口径总费用统计",
}

SUPPORTED_QUERY_KEYS = set(SUPPORTED_QUERY_KEY_DESCRIPTIONS.keys())


@dataclass
class QuestionRecord:
    """题库单题记录。

    说明：
        1. 统一保存题号、来源、原题和分层结论，便于后续继续迭代；
        2. 题库来源分为 230 / 73 / 600 三组，分别对应原物流题、真实问法、结构化题。
    """

    question_id: str
    question: str
    source_group: str
    source_label: str
    category_label: str
    difficulty: str
    classification: str
    reason: str
    query_key: str | None
    ability_boundary: str | None


class LogisticsQuestionBankClassifier:
    """物流域题库分层分类器。

    说明：
        1. 分类必须基于当前真实 data-qa 代码能力，而不是凭主观理解；
        2. 先按明确不支持 / 明确需澄清做规则兜底，再调用当前 planner 判断稳定支持边界；
        3. D 类只在发现“已有固定标准答案但仍需业务修订”的题时使用，本轮默认保留为空集。
    """

    UNSUPPORTED_KEYWORDS = (
        "预测",
        "预计",
        "未来",
        "趋势",
        "波动区间",
        "ETA",
        "到货时间",
        "到达时间",
        "风险评分",
        "评分模型",
        "根因分析",
        "原因分析",
        "方案设计",
        "设计一套",
        "优化策略",
        "调度策略",
    )

    META_QUESTION_KEYWORDS = (
        "应如何处理",
        "应遵循什么原则",
        "还能否直接输出",
        "应如何回答",
        "如何处理这种",
        "需先进行",
        "更稳妥的做法",
        "是否还能直接输出",
        "是否应该",
    )

    CLARIFICATION_KEYWORDS = (
        "最近",
        "最差",
        "异常",
        "有没有问题",
        "效率怎么样",
        "风险怎么样",
        "哪些有问题",
        "特殊订单",
        "是不是变高",
        "是否变高",
        "怎么样",
    )

    METRIC_KEYWORDS = (
        "发运量",
        "运量",
        "总费用",
        "总运费",
        "车次",
        "车辆数",
        "签收率",
        "元/瓦",
        "件数",
        "平均运费",
        "偏差率",
        "额外费用",
        "总发运量",
    )

    def __init__(self) -> None:
        self.planner = LogisticsDataQaPlanner()

    def classify(self, record: dict[str, str]) -> QuestionRecord:
        """对单题做 A/B/C/D 分层。"""
        question = record["question"].strip()
        compact = re.sub(r"\s+", "", question)

        if self._is_d_candidate(record, compact):
            return self._build_record(record, "D", "当前题存在固定标准答案待业务复核。", None)

        if self._is_explicit_unsupported(record, compact):
            return self._build_record(
                record,
                "C",
                "当前属于预测、开放讨论、治理原则或超出现有结构化数据问答范围的问题。",
                None,
            )

        plan = self.planner.build_plan(question)
        if plan.query_key in SUPPORTED_QUERY_KEYS and not plan.needs_clarification and plan.intent != "unsupported":
            return self._build_record(
                record,
                "A",
                f"当前已命中受控 query_key：{plan.query_key}，属于现有稳定支持范围。",
                plan.query_key,
            )

        if plan.intent == "unsupported":
            return self._build_record(
                record,
                "C",
                plan.unsupported_reason or "当前问题暂不支持。",
                plan.query_key,
            )

        if self._should_clarify(record, compact, plan.needs_clarification):
            return self._build_record(
                record,
                "B",
                "当前问题缺少关键时间、指标、异常定义或比较口径，应先澄清再执行查询。",
                plan.query_key,
            )

        return self._build_record(
            record,
            "C",
            "题意虽然明确，但当前物流数据问答主链路没有对应的稳定 query_key 支持。",
            plan.query_key,
        )

    def _is_explicit_unsupported(self, record: dict[str, str], compact: str) -> bool:
        """识别明确应归入不支持的问题。"""
        source_label = record["source_label"]
        category_label = record["category_label"]

        if "预测" in source_label or "预测" in category_label:
            return True

        if any(keyword in compact for keyword in self.UNSUPPORTED_KEYWORDS):
            return True

        if any(keyword in compact for keyword in self.META_QUESTION_KEYWORDS):
            return True

        if "额外费用" in compact and any(keyword in compact for keyword in ("项目", "原因", "明细")):
            return True

        if "吨" in compact:
            return True

        if "报价" in compact and "比较" in compact:
            return True

        return False

    def _should_clarify(self, record: dict[str, str], compact: str, planner_says_clarify: bool) -> bool:
        """识别应先澄清的问题。"""
        if any(keyword in compact for keyword in self.CLARIFICATION_KEYWORDS):
            return True

        has_year = bool(re.search(r"\d{2,4}年", compact))
        has_metric = any(keyword in compact for keyword in self.METRIC_KEYWORDS)

        if planner_says_clarify and (not has_year or not has_metric):
            return True

        if planner_says_clarify and any(keyword in compact for keyword in ("哪个", "哪些", "多少")) and not has_year:
            return True

        if record["source_group"] == "73" and planner_says_clarify:
            return True

        return False

    @staticmethod
    def _is_d_candidate(record: dict[str, str], compact: str) -> bool:
        """识别待修订标准答案的题。

        说明：
            当前题库中除已修订完成的 MVP 核心题外，没有更多“明确锁定答案但与真实 SQL 冲突”的题。
            因此本轮默认 D 类为空，只保留这个判断口以便后续继续扩展。
        """
        _ = record, compact
        return False

    @staticmethod
    def _build_record(
        record: dict[str, str],
        classification: str,
        reason: str,
        query_key: str | None,
    ) -> QuestionRecord:
        """组装结构化题库记录。"""
        return QuestionRecord(
            question_id=record["question_id"],
            question=record["question"],
            source_group=record["source_group"],
            source_label=record["source_label"],
            category_label=record["category_label"],
            difficulty=record["difficulty"],
            classification=classification,
            reason=reason,
            query_key=query_key,
            ability_boundary=SUPPORTED_QUERY_KEY_DESCRIPTIONS.get(query_key),
        )


def load_question_bank(docx_path: Path) -> list[dict[str, str]]:
    """从题库 docx 中读取三类题目。"""
    document = Document(str(docx_path))
    items: list[dict[str, str]] = []

    # 第一类：原 230 题
    for row in document.tables[1].rows[1:]:
        cells = [cell.text.strip() for cell in row.cells]
        items.append(
            {
                "question_id": cells[0],
                "source_group": "230",
                "source_label": cells[1],
                "category_label": cells[2],
                "difficulty": cells[3],
                "question": cells[4],
            }
        )

    # 第二类：原始业务问法 73 条
    for idx, row in enumerate(document.tables[2].rows[1:], start=1):
        cells = [cell.text.strip() for cell in row.cells]
        items.append(
            {
                "question_id": f"RAW{idx:03d}",
                "source_group": "73",
                "source_label": cells[1],
                "category_label": cells[1],
                "difficulty": "L-raw",
                "question": cells[2],
            }
        )

    # 第三类：新增结构化题 600 条
    for row in document.tables[3].rows[1:]:
        cells = [cell.text.strip() for cell in row.cells]
        items.append(
            {
                "question_id": cells[0],
                "source_group": "600",
                "source_label": cells[1],
                "category_label": cells[1],
                "difficulty": cells[2],
                "question": cells[3],
            }
        )

    return items


def build_summary(records: list[QuestionRecord]) -> dict[str, Any]:
    """生成分层统计摘要。"""
    counts = Counter(item.classification for item in records)
    by_source: dict[str, dict[str, int]] = defaultdict(lambda: defaultdict(int))
    by_query_key = Counter(item.query_key for item in records if item.query_key)
    examples: dict[str, list[dict[str, Any]]] = {"A": [], "B": [], "C": [], "D": []}

    for item in records:
        by_source[item.source_group][item.classification] += 1
        if len(examples[item.classification]) < 6:
            examples[item.classification].append(
                {
                    "question_id": item.question_id,
                    "question": item.question,
                    "source_group": item.source_group,
                    "reason": item.reason,
                    "query_key": item.query_key,
                }
            )

    return {
        "total_questions": len(records),
        "counts": {key: counts.get(key, 0) for key in ("A", "B", "C", "D")},
        "source_breakdown": {source: {key: bucket.get(key, 0) for key in ("A", "B", "C", "D")} for source, bucket in by_source.items()},
        "query_key_breakdown": dict(sorted(by_query_key.items())),
        "examples": examples,
        "stable_support_boundary": list(SUPPORTED_QUERY_DESCRIPTIONS_FOR_DOC()),
        "clarification_boundary": [
            "问题缺少时间范围，例如只问“历史总费用是多少”但没有年份或来源范围。",
            "问题缺少指标口径，例如“最近物流成本是不是变高了”。",
            "问题缺少比较维度、异常定义或评价标准，例如“哪个最差”“哪些有问题”。",
        ],
        "unsupported_boundary": [
            "预测、预计、未来趋势、波动区间类问题。",
            "ETA、到货时间、复杂时效推理类问题。",
            "开放讨论、治理原则、如何设计、如何处理极端值等非结构化问答。",
            "当前 MVP 明确不支持的额外费用项目/原因/明细类问题。",
        ],
        "standard_answer_revision_candidates": [],
    }


def SUPPORTED_QUERY_DESCRIPTIONS_FOR_DOC() -> list[str]:
    """返回文档里可直接展示的稳定支持边界。"""
    return [
        "按区域汇总历史发运件数、历史车次。",
        "按区域统计运输方式平均元/瓦，并做成本排序。",
        "按省份统计城市总费用排名前五。",
        "按始发地/省份/车型统计月均运费。",
        "按年份统计额外费用占比最高月份。",
        "按始发地与指定承运商统计总运费。",
        "按客户/项目统计历史发运量 MW。",
        "按车型统计历史车次。",
        "按最终客户统计同一客户多个始发地发货情况。",
        "按区域统计计划件数与实际件数偏差率。",
        "2026 年按 pickup_date 统计发运量 MW 与车次。",
        "2026 年承运商 SIGNEDFOR 签收率前十/后十。",
        "2026 特殊业务口径总费用（经营计划、辅料送样、刘娟用车）。",
        "2026 已建档但无任务承运商识别。",
    ]


def render_markdown(summary: dict[str, Any], records: list[QuestionRecord]) -> str:
    """生成正式题库分层结论文档。"""
    examples = summary["examples"]
    counts = summary["counts"]
    lines: list[str] = []
    lines.append("# 物流域题库分层验证与收口")
    lines.append("")
    lines.append("## 结论")
    lines.append("")
    lines.append(
        f"当前物流域题库总量为 **{summary['total_questions']}** 条。"
        f"基于当前真实物流数据问答主链路能力，当前分层结果为："
        f"A 类 {counts['A']} 条、B 类 {counts['B']} 条、C 类 {counts['C']} 条、D 类 {counts['D']} 条。"
    )
    lines.append("")
    lines.append("这份结果用于明确：")
    lines.append("- 物流数据问答 MVP 已收口，核心能力可稳定支撑一批固定结构化题。")
    lines.append("- 但物流域全量题库 **尚未完全收口**，不能把 20/20 的 MVP 验收误判成 903 条题库都已稳定通过。")
    lines.append("")
    lines.append("## A / B / C / D 定义")
    lines.append("")
    lines.append("- A 类：当前已支持，应通过。题意明确、口径已锁定，且当前主链路已有稳定 query_key。")
    lines.append("- B 类：当前应澄清。问题缺少时间、指标、评价标准或异常定义，不应直接给结果。")
    lines.append("- C 类：当前应明确不支持。超出当前 MVP 边界，例如预测、ETA、开放讨论或复杂设计问题。")
    lines.append("- D 类：当前需修订标准答案或业务口径。本轮新增题库中未发现新的明确 D 类题，当前为 0。")
    lines.append("")
    lines.append("## 分层统计")
    lines.append("")
    lines.append("| 分类 | 数量 |")
    lines.append("| --- | ---: |")
    for key in ("A", "B", "C", "D"):
        lines.append(f"| {key} 类 | {counts[key]} |")
    lines.append("")
    lines.append("### 按来源分类统计")
    lines.append("")
    lines.append("| 来源 | A | B | C | D |")
    lines.append("| --- | ---: | ---: | ---: | ---: |")
    for source in ("230", "73", "600"):
        bucket = summary["source_breakdown"].get(source, {})
        lines.append(
            f"| {source} | {bucket.get('A', 0)} | {bucket.get('B', 0)} | {bucket.get('C', 0)} | {bucket.get('D', 0)} |"
        )
    lines.append("")
    lines.append("## 当前真正稳定支持的题型边界")
    lines.append("")
    for item in summary["stable_support_boundary"]:
        lines.append(f"- {item}")
    lines.append("")
    lines.append("## 当前应先澄清的典型模式")
    lines.append("")
    for item in summary["clarification_boundary"]:
        lines.append(f"- {item}")
    lines.append("")
    lines.append("## 当前应明确不支持的题型边界")
    lines.append("")
    for item in summary["unsupported_boundary"]:
        lines.append(f"- {item}")
    lines.append("")

    for label, title in (("A", "A 类代表样例"), ("B", "B 类代表样例"), ("C", "C 类代表样例"), ("D", "D 类代表样例")):
        lines.append(f"## {title}")
        lines.append("")
        if not examples[label]:
            lines.append("- 当前没有新增样例。")
            lines.append("")
            continue
        for item in examples[label]:
            lines.append(
                f"- {item['question_id']}（来源 {item['source_group']}）：{item['question']}"
                f"。原因：{item['reason']}"
                + (f" 关联 query_key：`{item['query_key']}`。" if item["query_key"] else "")
            )
        lines.append("")

    lines.append("## 当前需修订标准答案 / 业务口径的题")
    lines.append("")
    if not summary["standard_answer_revision_candidates"]:
        lines.append("- 本轮题库分层中，未发现新的明确 D 类题。")
        lines.append("- 此前 MVP 核心题中的 Q02 / Q17 / Q19 等争议项已在前序主线中完成基线修订，不再重复计入本轮 D 类。")
    else:
        for item in summary["standard_answer_revision_candidates"]:
            lines.append(f"- {item}")
    lines.append("")
    lines.append("## 结论建议")
    lines.append("")
    lines.append("- 可以明确说：**物流数据问答 MVP 已收口。**")
    lines.append("- 也必须明确说：**物流域全量题库尚未完全收口。**")
    lines.append("- 下一步最合理的动作，不是直接进入新域开发，而是：")
    lines.append("  1. 先针对 A 类题做批量自动回归；")
    lines.append("  2. 再把 B / C 类题的前端文案和响应策略固化；")
    lines.append("  3. 最后再决定是否扩下一批 query_key。")
    lines.append("")

    return "\n".join(lines)


def main() -> None:
    parser = argparse.ArgumentParser(description="物流域题库分层验证与收口")
    parser.add_argument(
        "--docx",
        default="/Users/zhuchangchao/Desktop/01_工作业务/计划经营部/项目/经营计划智能助手/一期/正式开始搭代码前需要提供的资料/2、真实问题集/物流域提问问题集.docx",
        help="题库 docx 路径",
    )
    parser.add_argument(
        "--json-output",
        default="/Users/zhuchangchao/Work/PythonProject/project/gcl-bp-ai/tmp/logistics_question_bank/logistics_question_bank_classification.json",
        help="分类 JSON 输出路径",
    )
    parser.add_argument(
        "--doc-output",
        default="/Users/zhuchangchao/Work/PythonProject/project/gcl-bp-ai/docs/LOGISTICS_QUESTION_BANK_CLASSIFICATION.md",
        help="分类文档输出路径",
    )
    args = parser.parse_args()

    docx_path = Path(args.docx)
    json_output = Path(args.json_output)
    doc_output = Path(args.doc_output)

    items = load_question_bank(docx_path)
    classifier = LogisticsQuestionBankClassifier()
    records = [classifier.classify(item) for item in items]
    summary = build_summary(records)

    json_output.parent.mkdir(parents=True, exist_ok=True)
    doc_output.parent.mkdir(parents=True, exist_ok=True)

    payload = {
        "docx_path": str(docx_path),
        "summary": summary,
        "items": [asdict(item) for item in records],
    }
    json_output.write_text(json.dumps(payload, ensure_ascii=False, indent=2), encoding="utf-8")
    doc_output.write_text(render_markdown(summary, records), encoding="utf-8")

    print(json.dumps(summary, ensure_ascii=False, indent=2))


if __name__ == "__main__":
    main()
