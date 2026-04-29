from __future__ import annotations

import json
import re
import shutil
import sys
import zipfile
from pathlib import Path
from typing import Any

import pandas as pd

REPO_ROOT = Path(__file__).resolve().parents[1]
if str(REPO_ROOT) not in sys.path:
    sys.path.insert(0, str(REPO_ROOT))

from docx import Document
from sqlalchemy import create_engine
from sqlalchemy.orm import Session, sessionmaker

from backend.app.db.base import Base
from backend.app.domains.plan_bom.models import PlanBomHeader, PlanBomMaterialLine
from backend.app.domains.plan_bom.repositories.import_repository import PlanBomImportRepository
from backend.app.domains.plan_bom.repositories.query_repository import PlanBomQueryRepository
from backend.app.domains.plan_bom.services.answer_presentation_service import PlanBomAnswerPresentationService
from backend.app.domains.plan_bom.services.excel_import_service import PlanBomExcelImportService
from backend.app.domains.plan_bom.services.nlu_center_service import PlanBomNluCenterService
from backend.app.domains.plan_bom.services.qa_service import PlanBomQaService
from backend.app.domains.plan_bom.services.query_service import PlanBomQueryService
from backend.app.models.sys_query_log import SysQueryLog  # noqa: F401


TMP_DIR = REPO_ROOT / "tmp" / "plan_bom"
CONFIG_DIR = REPO_ROOT / "backend" / "app" / "domains" / "plan_bom" / "config"
INPUT_DIR = TMP_DIR / "input"
DEFAULT_SOURCE_ZIP_CANDIDATES = [
    REPO_ROOT / "data" / "plan_bom" / "BOM 源数据.zip",
    INPUT_DIR / "BOM 源数据.zip",
]
DEFAULT_QUESTION_FILE_CANDIDATES = [
    REPO_ROOT / "data" / "plan_bom" / "BOM问题.xlsx",
    INPUT_DIR / "BOM问题.xlsx",
    REPO_ROOT / "tmp" / "plan_bom_medium_sample" / "BOM问题.xlsx",
]
QUESTION_COLUMN_CANDIDATES = ("问题文本", "问题", "提问", "用户问题", "问题内容", "question")
RUNTIME_DB = TMP_DIR / "plan_bom_runtime.db"


def ensure_dirs() -> None:
    """创建计划 BOM 脚本输出目录。

    返回：
        无返回值。
    """

    TMP_DIR.mkdir(parents=True, exist_ok=True)
    CONFIG_DIR.mkdir(parents=True, exist_ok=True)
    INPUT_DIR.mkdir(parents=True, exist_ok=True)


def resolve_existing_path(explicit_path: str | Path | None, candidates: list[Path], label: str) -> Path:
    """解析输入文件路径，避免把本机路径写死到脚本中。

    参数：
        explicit_path: 命令行显式传入的文件路径；
        candidates: 项目内约定默认候选路径；
        label: 报错展示的业务文件名称。

    返回：
        已存在的文件路径；不存在时抛出清晰错误。
    """

    ensure_dirs()
    if explicit_path:
        path = Path(explicit_path).expanduser().resolve()
        if not path.exists():
            raise FileNotFoundError(f"{label}不存在：{path}")
        return path
    for candidate in candidates:
        if candidate.exists():
            return candidate.resolve()
    candidate_text = "；".join(str(item) for item in candidates)
    raise FileNotFoundError(f"{label}缺失。请通过参数指定，默认查找路径：{candidate_text}")


def build_runtime_session(*, reset: bool = False) -> Session:
    """创建或打开计划 BOM 本地 SQLite 运行库。

    参数：
        reset: 是否删除旧库后重建。

    返回：
        SQLAlchemy Session。
    """

    ensure_dirs()
    if reset and RUNTIME_DB.exists():
        RUNTIME_DB.unlink()
    engine = create_engine(f"sqlite:///{RUNTIME_DB}")
    Base.metadata.create_all(engine)
    return sessionmaker(bind=engine)()


def extract_bom_zip(source_zip: str | Path | None = None) -> list[Path]:
    """解压 BOM 源数据并过滤无效文件。

    参数：
        source_zip: 用户提供的 BOM 源数据 zip。

    返回：
        有效 Excel 文件路径列表。
    """

    source_path = resolve_existing_path(source_zip, DEFAULT_SOURCE_ZIP_CANDIDATES, "BOM 源数据 zip")
    extract_dir = TMP_DIR / "source_unzipped"
    if extract_dir.exists():
        shutil.rmtree(extract_dir)
    extract_dir.mkdir(parents=True, exist_ok=True)
    with zipfile.ZipFile(source_path) as zf:
        for info in zf.infolist():
            name = info.filename
            if "__MACOSX" in name or Path(name).name.startswith("._") or name.endswith("/"):
                continue
            suffix = Path(name).suffix.lower()
            if suffix not in {".xls", ".xlsx", ".xlsm"}:
                continue
            zf.extract(info, extract_dir)
    files = sorted([path for path in extract_dir.rglob("*") if path.suffix.lower() in {".xls", ".xlsx", ".xlsm"}])
    if not files:
        raise RuntimeError("BOM 源数据 zip 中没有可解析的 Excel 文件。")
    return files


def import_source_zip(*, source_zip: str | Path | None = None, reset: bool = True) -> dict[str, Any]:
    """导入真实 BOM 源数据到本地运行库。

    参数：
        reset: 是否重建本地运行库。

    返回：
        导入汇总报告。
    """

    source_path = resolve_existing_path(source_zip, DEFAULT_SOURCE_ZIP_CANDIDATES, "BOM 源数据 zip")
    files = extract_bom_zip(source_path)
    session = build_runtime_session(reset=reset)
    service = PlanBomExcelImportService(repository=PlanBomImportRepository(session))
    reports = []
    for path in files:
        report = service.import_file(path)
        reports.append(report.model_dump(mode="json"))
    summary = {
        "source_zip": str(source_path),
        "runtime_db": str(RUNTIME_DB),
        "file_count": len(files),
        "success_count": sum(1 for item in reports if item["status"] == "success"),
        "failed_count": sum(1 for item in reports if item["status"] != "success"),
        "parsed_orders_count": sum(item["header_count"] for item in reports),
        "parsed_materials_count": sum(item["material_line_count"] for item in reports),
        "warning_count": sum(item["warning_count"] for item in reports),
        "error_count": sum(item["error_count"] for item in reports),
        "reports": reports,
    }
    (TMP_DIR / "plan_bom_source_import_report.json").write_text(json.dumps(summary, ensure_ascii=False, indent=2), encoding="utf-8")
    return summary


def make_qa_service(session: Session | None = None, *, presentation_enabled: bool = False) -> PlanBomQaService:
    """创建基于本地运行库的 BOM QA 服务。

    参数：
        session: 可选数据库会话；
        presentation_enabled: 是否启用 live LLM 表达层，默认关闭以保证常规回归稳定。

    返回：
        PlanBomQaService 实例。
    """

    actual_session = session or build_runtime_session(reset=False)
    repository = PlanBomQueryRepository(actual_session)
    return PlanBomQaService(
        repository=repository,
        query_service=PlanBomQueryService(repository=repository),
        nlu_service=PlanBomNluCenterService(repository=repository),
        presentation_service=PlanBomAnswerPresentationService(enabled=presentation_enabled),
    )


def build_standardized_outputs(session: Session | None = None) -> dict[str, Any]:
    """生成标准化 BOM 数据、订单索引和材料别名配置。

    参数：
        session: 可选数据库会话。

    返回：
        标准化输出汇总。
    """

    actual_session = session or build_runtime_session(reset=False)
    headers = actual_session.query(PlanBomHeader).filter(PlanBomHeader.is_active == 1).all()
    lines = actual_session.query(PlanBomMaterialLine).all()
    standardized_lines = [
        {
            "order_no": line.order_no,
            "version_no": line.version_no,
            "order_identity_key": line.order_identity_key,
            "file_instance_key": line.file_instance_key,
            "sap_code": line.sap_code,
            "material_name": line.material_name,
            "material_category": line.material_category,
            "description": line.description,
            "standard_usage": str(line.standard_usage) if line.standard_usage is not None else None,
            "unit": line.unit,
            "remark": line.remark,
            "source_file": line.import_batch_id,
            "raw_row_no": line.raw_row_no,
        }
        for line in lines
    ]
    order_index = []
    for header in headers:
        aliases = sorted(
            {
                header.order_no,
                header.order_no[-5:] if len(header.order_no) >= 5 else header.order_no,
                f"{header.order_no[-10:]}" if len(header.order_no) >= 10 else header.order_no,
                *(re.findall(r"NT[0-9A-Z]+/[0-9A-Z]+GDF|NT[0-9A-Z]+GDF", header.order_name or "", flags=re.I)),
            }
        )
        order_index.append(
            {
                "order_no": header.order_no,
                "order_name": header.order_name,
                "version_no": header.version_no,
                "order_identity_key": header.order_identity_key,
                "file_instance_key": header.file_instance_key,
                "effective_date": str(header.effective_date) if header.effective_date else None,
                "raw_file_name": header.raw_file_name,
                "aliases": [item for item in aliases if item],
            }
        )
    material_aliases = json.loads((CONFIG_DIR / "material_aliases.json").read_text(encoding="utf-8"))
    (CONFIG_DIR / "plan_bom_order_index.json").write_text(json.dumps(order_index, ensure_ascii=False, indent=2), encoding="utf-8")
    (TMP_DIR / "plan_bom_standardized_materials.json").write_text(json.dumps(standardized_lines, ensure_ascii=False, indent=2), encoding="utf-8")
    return {
        "orders": len(headers),
        "materials": len(lines),
        "order_index_path": str(CONFIG_DIR / "plan_bom_order_index.json"),
        "standardized_data_path": str(TMP_DIR / "plan_bom_standardized_materials.json"),
        "material_aliases": material_aliases,
    }


def read_question_file(question_file: str | Path | None = None) -> tuple[list[dict[str, Any]], dict[str, Any]]:
    """读取 BOM 问题文件，正式源优先为 BOM问题.xlsx。

    参数：
        question_file: 用户提供的问题文件，支持 xlsx/xls/docx。

    返回：
        二元组：(问题记录列表, 文件元信息)。
    """

    path = resolve_existing_path(question_file, DEFAULT_QUESTION_FILE_CANDIDATES, "BOM 问题文件")
    suffix = path.suffix.lower()
    if suffix in {".xlsx", ".xls"}:
        records, metadata = read_question_xlsx(path)
    elif suffix == ".docx":
        records, metadata = read_question_docx(path)
    else:
        raise ValueError(f"BOM 问题文件格式不支持：{path.suffix}，仅支持 .xlsx/.xls/.docx")
    metadata.update(
        {
            "question_file": str(path),
            "question_file_name": path.name,
            "question_file_type": suffix.lstrip("."),
            "question_count": len(records),
            "formal_source": suffix in {".xlsx", ".xls"},
            "docx_compatibility_only": suffix == ".docx",
        }
    )
    return records, metadata


def read_question_xlsx(question_xlsx: Path) -> tuple[list[dict[str, Any]], dict[str, Any]]:
    """读取正式 BOM 问题 xlsx 并容错识别表头。

    参数：
        question_xlsx: 用户提供的 BOM问题.xlsx。

    返回：
        二元组：(问题记录列表, 文件元信息)。
    """

    workbook = pd.ExcelFile(question_xlsx)
    selected_sheet = ""
    selected_frame: pd.DataFrame | None = None
    selected_question_column = ""
    sheet_summaries: list[dict[str, Any]] = []
    for sheet_name in workbook.sheet_names:
        frame = pd.read_excel(workbook, sheet_name=sheet_name, dtype=object)
        frame = frame.dropna(how="all")
        question_column = _detect_question_column(frame)
        valid_count = int(frame[question_column].notna().sum()) if question_column else 0
        sheet_summaries.append(
            {
                "sheet": sheet_name,
                "rows": int(len(frame)),
                "question_column": question_column,
                "valid_question_rows": valid_count,
            }
        )
        current_best = int(selected_frame[selected_question_column].notna().sum()) if selected_frame is not None else -1
        if question_column and valid_count > current_best:
            selected_sheet = sheet_name
            selected_frame = frame
            selected_question_column = question_column
    if selected_frame is None or not selected_question_column:
        raise ValueError(f"无法在 {question_xlsx} 中识别问题列，请检查是否包含“问题文本/问题/提问”等列。")

    records: list[dict[str, Any]] = []
    skipped_rows: list[dict[str, Any]] = []
    for row_index, row in selected_frame.iterrows():
        text = _clean_cell(row.get(selected_question_column))
        if not _is_valid_question_text(text):
            skipped_rows.append({"row_no": int(row_index) + 2, "reason": "问题文本为空或不是有效问题"})
            continue
        serial = _clean_cell(_get_first_existing(row, ("序号", "编号", "id", "ID"))) or f"Q{len(records) + 1:03d}"
        source = _clean_cell(_get_first_existing(row, ("来源", "source"))) or f"xlsx:{selected_sheet}"
        category = _clean_cell(_get_first_existing(row, ("问题类别", "类别", "分类", "intent"))) or "未分类"
        role = _clean_cell(_get_first_existing(row, ("适用角色", "角色"))) or "计划/BOM业务人员"
        scenario = _clean_cell(_get_first_existing(row, ("业务场景", "场景"))) or "计划 BOM 问答"
        data_context = _clean_cell(_get_first_existing(row, ("对应业务数据内容", "业务数据内容", "数据内容", "标准答案", "备注"))) or "来源于用户提供的 BOM 源数据和 BOM问题.xlsx。"
        raw_fields = {str(column): _clean_cell(value) for column, value in row.items() if _clean_cell(value)}
        records.append(
            {
                "序号": str(serial),
                "来源": source,
                "问题类别": category,
                "适用角色": role,
                "业务场景": scenario,
                "问题文本": text,
                "对应业务数据内容": data_context,
                "原始字段": raw_fields,
                "问题来源文件": question_xlsx.name,
                "问题来源sheet": selected_sheet,
                "问题来源行号": int(row_index) + 2,
            }
        )
    return records, {
        "selected_sheet": selected_sheet,
        "selected_question_column": selected_question_column,
        "sheet_summaries": sheet_summaries,
        "skipped_rows": skipped_rows,
    }


def read_question_docx(question_docx: Path) -> tuple[list[dict[str, Any]], dict[str, Any]]:
    """兼容读取旧 DOCX 问题源，当前不作为正式问题来源。

    参数：
        question_docx: 兼容输入的问题样例文档。

    返回：
        二元组：(问题记录列表, 文件元信息)。
    """

    if not question_docx.exists():
        raise FileNotFoundError(f"BOM 问题样例缺失：{question_docx}")
    doc = Document(question_docx)
    questions: list[dict[str, Any]] = []
    index = 1
    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue
        match = re.match(r"^\d+[、.]\s*(.+)", text)
        if match:
            questions.append(_question_record(index, match.group(1), "DOCX段落", "原始问题"))
            index += 1
        elif "使用功率预测" in text:
            questions.append(_question_record(index, text, "DOCX段落", "真实问题"))
            index += 1
    for table in doc.tables:
        for row in table.rows:
            text = " ".join(cell.text.strip() for cell in row.cells if cell.text.strip()).strip()
            if text:
                questions.append(_question_record(index, text, "DOCX表格", "真实问题"))
                index += 1
    return questions, {"selected_sheet": None, "selected_question_column": None, "sheet_summaries": [], "skipped_rows": []}


def _detect_question_column(frame: pd.DataFrame) -> str:
    """从 xlsx 表头中容错识别问题列。

    参数：
        frame: 单个 sheet 的 DataFrame。

    返回：
        问题列名；识别失败返回空字符串。
    """

    columns = [str(column).strip() for column in frame.columns]
    for candidate in QUESTION_COLUMN_CANDIDATES:
        for column in columns:
            if candidate.lower() == column.lower():
                return column
    for column in columns:
        if any(keyword in column for keyword in ("问题", "提问", "question")):
            return column
    scored: list[tuple[int, str]] = []
    for column in columns:
        sample = frame[column].dropna().astype(str).head(50).tolist()
        score = sum(1 for item in sample if _is_valid_question_text(item))
        if score:
            scored.append((score, column))
    return sorted(scored, reverse=True)[0][1] if scored else ""


def _get_first_existing(row: pd.Series, columns: tuple[str, ...]) -> Any:
    """从一行中读取首个存在且非空的字段。

    参数：
        row: DataFrame 行；
        columns: 候选列名。

    返回：
        单元格值；不存在时返回 None。
    """

    for column in columns:
        if column in row and _clean_cell(row.get(column)):
            return row.get(column)
    return None


def _clean_cell(value: Any) -> str:
    """清洗 Excel/Word 单元格文本。

    参数：
        value: 原始单元格值。

    返回：
        去除空白和 nan 后的字符串。
    """

    if value is None:
        return ""
    text = str(value).strip()
    return "" if text.lower() in {"nan", "none", "nat"} else text


def _is_valid_question_text(text: str) -> bool:
    """判断一行文本是否可作为问题进入正式回归。

    参数：
        text: 待判断文本。

    返回：
        有效问题返回 True。
    """

    normalized = _clean_cell(text)
    if len(normalized) < 4:
        return False
    if normalized.startswith("._") or normalized.startswith("__MACOSX"):
        return False
    return bool(
        any(
            mark in normalized
            for mark in ("?", "？", "哪些", "是否", "有没有", "怎么", "什么", "查询", "列出", "对比", "比较", "生成", "规格", "物料", "BOM", "bom")
        )
    )


def _question_record(index: int, text: str, source: str, category: str) -> dict[str, Any]:
    """构造问题台账基础记录。

    参数：
        index: 序号；
        text: 问题文本；
        source: 来源；
        category: 问题类别。

    返回：
        问题记录。
    """

    return {
        "序号": f"Q{index:03d}",
        "来源": source,
        "问题类别": category,
        "适用角色": "计划/BOM业务人员",
        "业务场景": "计划 BOM 问答",
        "问题文本": text,
        "对应业务数据内容": "来源于用户提供的 BOM 源数据和问题样例。",
    }


def generate_variants(question: str) -> list[str]:
    """为问题生成两条语义等价变体。

    参数：
        question: 原始问题。

    返回：
        两条变体问法。
    """

    core = question.strip().rstrip("？?")
    variant1 = f"请按计划 BOM 数据核对：{core}，需要返回可追溯的订单和材料信息。"
    variant2 = f"从业务核对角度看，{core}；请说明能直接回答、需要追问还是当前无法回答。"
    return [variant1, variant2]


def write_markdown(path: Path, title: str, lines: list[str]) -> None:
    """写 Markdown 文档。

    参数：
        path: 输出路径；
        title: 文档标题；
        lines: 正文行。

    返回：
        无返回值。
    """

    path.write_text("\n".join([f"# {title}", "", *lines, ""]) , encoding="utf-8")


__all__ = [
    "CONFIG_DIR",
    "DEFAULT_QUESTION_FILE_CANDIDATES",
    "DEFAULT_SOURCE_ZIP_CANDIDATES",
    "INPUT_DIR",
    "RUNTIME_DB",
    "TMP_DIR",
    "build_runtime_session",
    "build_standardized_outputs",
    "ensure_dirs",
    "extract_bom_zip",
    "generate_variants",
    "import_source_zip",
    "make_qa_service",
    "read_question_file",
    "read_question_docx",
    "read_question_xlsx",
    "resolve_existing_path",
    "write_markdown",
]
